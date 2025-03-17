from docx import Document
from docx.shared import Inches
import os
from natsort import natsorted  # Natural sorting

def png_to_word(folder_path, output_docx):
    # Get all PNG files in the folder (case-insensitive)
    png_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.png')]

    # Sort the files in natural numerical order
    png_files = natsorted(png_files)

    if not png_files:
        print("No PNG files found in the folder.")
        return

    # Create a new Word document
    doc = Document()

    for png in png_files:
        img_path = os.path.join(folder_path, png)
        
        # Add image to the document
        doc.add_picture(img_path, width=Inches(5))  # Adjust width as needed
        doc.add_paragraph(png)  # Add the filename as a caption

    # Save the document
    doc.save(output_docx)
    print(f"Word document saved as {output_docx}")

# Corrected file paths with raw string notation
folder_path = r"folder_path"
output_docx = r"output_folder_path\output.doc"

png_to_word(folder_path, output_docx)
